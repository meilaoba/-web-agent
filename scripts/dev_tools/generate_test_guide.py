"""生成《RAG 知识库功能测试指南》Word 文档。

产出: docs/RAG知识库功能测试指南.docx

用法:
    python scripts/dev_tools/generate_test_guide.py

依赖: python-docx（pip install python-docx）
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT = PROJECT_ROOT / "docs" / "RAG知识库功能测试指南.docx"

CODE_FONT = "Consolas"
BODY_FONT = "宋体"
BODY_FONT_EN = "Calibri"


# ---------- 样式辅助 ----------
def set_run_font(run, name_cn=BODY_FONT, name_en=BODY_FONT_EN, size=10.5, bold=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置东亚字体，避免中文回退
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name_cn)


def add_para(doc, text, size=10.5, bold=False, color=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, size=10.5, indent=0.5, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p


def add_code_block(doc, code: str, size=9):
    """等宽字体代码块，浅灰底纹。"""
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(line if line else " ")
        set_run_font(run, name_cn="等线", name_en=CODE_FONT, size=size)
        # 段落底纹
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
    # 代码块后留空
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return sp


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=9.5, bold=True)
        # 表头底纹
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tc_pr.append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        set_run_font(run, size=16, bold=True, color=(0x1F, 0x3B, 0x63))
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        set_run_font(run, size=13, bold=True, color=(0x2E, 0x5B, 0x9E))
    return h


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        set_run_font(run, size=11, bold=True, color=(0x33, 0x33, 0x33))
    return h


# ---------- 文档内容 ----------
def build() -> Path:
    doc = Document()

    # 标题
    title = doc.add_heading("RAG 知识库第一阶段 功能测试指南", level=0)
    for run in title.runs:
        set_run_font(run, size=22, bold=True, color=(0x1F, 0x3B, 0x63))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "项目：AI驱动的Web代码安全审计多Agent系统", size=11, bold=True)
    add_para(doc, "适用范围：RAG 知识库 Phase 1 数据链路（文档加载 / 文本清洗 / 文本分割 / Metadata / 处理管线）", size=11)
    add_para(doc, "更新日期：2026-08-17", size=11, space_after=12)

    # ================= 1. 概述 =================
    add_h1(doc, "1. 概述")
    add_para(doc, "本指南说明如何验证 RAG 知识库第一阶段的各个功能模块是否正常。测试分为两种方式：")
    add_bullet(doc, "自动化测试：pytest 单元测试，一条命令运行全部 50 个用例，覆盖各模块与端到端链路；")
    add_bullet(doc, "手动交互测试：使用 Python 交互式命令逐模块验证，便于理解每个功能的行为与边界。")
    add_para(doc, "测试数据：")
    add_bullet(doc, "backend/tests/fixtures/ —— 5 个格式的少量样例（md / txt / json / html / pdf），用于单元测试；")
    add_bullet(doc, "backend/data/raw/ —— 正式知识数据 16 个文件（10 类漏洞 + OWASP Top 10 + 多格式参考），用于完整管线验证。")

    # ================= 2. 环境准备 =================
    add_h1(doc, "2. 测试环境准备")
    add_h2(doc, "2.1 安装依赖")
    add_para(doc, "要求 Python 3.10 及以上版本。在项目根目录执行：")
    add_code_block(doc, """pip install -r requirements.txt        # 运行依赖（python-dotenv / pypdf / beautifulsoup4）
pip install -r requirements-dev.txt    # 测试依赖（pytest / reportlab）""")
    add_para(doc, "说明：本开发环境因权限限制，依赖安装到项目内 .deps 目录并通过 PYTHONPATH 使用；普通环境直接 pip install 即可。项目提供两个一键脚本自动处理：")
    add_code_block(doc, """scripts\dev_tools\run_tests.cmd      # 一键运行全部测试
scripts\data_pipeline\run_pipeline.cmd  # 一键运行知识库处理管线""")

    add_h2(doc, "2.2 配置文件（可选）")
    add_para(doc, "复制 .env.example 为 .env 后可调整 Chunk 参数与日志级别；不配置时使用默认值（chunk_size=800，chunk_overlap=100），测试不受影响。")

    # ================= 3. 自动化测试 =================
    add_h1(doc, "3. 方式一：自动化测试（pytest）")
    add_h2(doc, "3.1 运行命令")
    add_para(doc, "一键运行（Windows）：")
    add_code_block(doc, "scripts\\dev_tools\\run_tests.cmd -v")
    add_para(doc, "或手动运行：")
    add_code_block(doc, """cd backend
python -m pytest tests -v""")
    add_para(doc, "运行单个测试文件 / 单个用例：")
    add_code_block(doc, """python -m pytest tests/test_loader.py -v          # 只测加载模块
python -m pytest tests/test_splitter.py::TestSplitBasics -v   # 只测某个测试类""")

    add_h2(doc, "3.2 预期结果")
    add_para(doc, "全部通过时输出：50 passed。若有失败用例，按第 6 节排查。")

    add_h2(doc, "3.3 测试文件与覆盖范围")
    add_table(
        doc,
        ["测试文件", "覆盖模块", "主要验证点"],
        [
            ["tests/test_loader.py", "文档加载", "5 种格式均可加载；注册表自动分发；metadata（source/document_name/file_type/category）完整；md 标题与 CWE 提取；json 白名单键提升；html 导航剔除；pdf 页数记录；目录批量加载与扩展名过滤"],
            ["tests/test_cleaner.py", "文本清洗", "噪声行删除（导航/版权）；连续重复行去重；控制字符/零宽字符清理；代码块内容完整保留；Payload（<script>）不被误删；CWE 编号保留；行首缩进保留；metadata 不丢失"],
            ["tests/test_splitter.py", "文本分割", "长文档生成多个 Chunk；Chunk token 在合理范围；内容不丢失；代码块不被从中切开；section 语义信息；chunk_id 唯一；基础 metadata 保留并扩展 chunk_index/chunk_id/token_count"],
            ["tests/test_metadata.py", "Metadata", "CWE/CVE/OWASP 编号提取；规范化后值类型与 ChromaDB 兼容；None/空值剔除；复杂对象降级；增强不覆盖已有值"],
            ["tests/test_pipeline.py", "端到端管线", "load→clean→enrich→split→normalize 全链路；产物可落盘且回读正确；每个 Chunk 保留来源溯源信息"],
        ],
        col_widths=[4.2, 2.6, 9.0],
    )
    add_para(doc, "", space_after=4)

    # ================= 4. 手动测试 =================
    add_h1(doc, "4. 方式二：各功能手动测试")
    add_para(doc, "以下命令请在项目根目录执行。先设置 Python 环境（一键脚本已封装；手动执行时需保证依赖可导入）。")

    # ---- 4.1 文档加载 ----
    add_h2(doc, "4.1 文档加载测试")
    add_h3(doc, "测试目的")
    add_para(doc, "验证各格式知识文件能被正确加载为统一 Document 结构，且 metadata 完整。")
    add_h3(doc, "操作步骤")
    add_para(doc, "（1）加载单个文件：")
    add_code_block(
        doc,
        "cd backend\n"
        'python -c "from app.rag.loader import load_documents; '
        "d=load_documents('data/raw/owasp_cwe/sql_injection.md')[0]; "
        'print(d.metadata); print(d.page_content[:80])"',
    )
    add_para(doc, "（2）加载整个知识目录（含递归）：")
    add_code_block(
        doc,
        'python -c "from app.rag.loader import load_documents_from_dir; '
        "ds=load_documents_from_dir('data/raw'); "
        "print('文档数:', len(ds))\"",
    )
    add_para(doc, "（3）按扩展名过滤：")
    add_code_block(
        doc,
        'python -c "from app.rag.loader import load_documents_from_dir; '
        "ds=load_documents_from_dir('data/raw', extensions=['.md','.json']); "
        "print(sorted({d.metadata['file_type'] for d in ds}))\"",
    )
    add_h3(doc, "预期结果与验证点")
    add_table(
        doc,
        ["验证点", "预期"],
        [
            ["目录加载", "打印文档数 = 16，且加载日志显示失败数 0"],
            ["md 加载", "metadata 含 title（一级标题）、cwe_id 列表（如 CWE-89）；正文含代码示例"],
            ["json 加载", "cwe_id / vulnerability_type / severity 等白名单键提升为 metadata，其余字段进入正文"],
            ["html 加载", "正文中不含首页/关于我们/版权所有等导航页脚噪声；metadata.title 为页面标题"],
            ["pdf 加载", "正文含文本内容；metadata 含 page_count"],
            ["txt 加载", "正文完整，metadata.file_type = txt"],
            ["所有格式", "metadata 均包含 source（绝对路径）/ document_name / file_type / category"],
            ["异常处理", "加载不存在的文件或未注册格式（如 .docx）时抛出 DocumentLoadError 且不崩溃"],
        ],
        col_widths=[4.0, 11.8],
    )
    add_para(doc, "", space_after=4)

    # ---- 4.2 文本清洗 ----
    add_h2(doc, "4.2 文本清洗测试")
    add_h3(doc, "测试目的")
    add_para(doc, "验证清洗器能删除明显噪声，且不破坏安全知识（代码示例 / Payload / CWE 编号）。")
    add_h3(doc, "操作步骤")
    add_para(doc, "（1）删除噪声验证：")
    add_code_block(doc, """python -c "
from app.rag.cleaner import TextCleaner
t = '首页\\n## 漏洞描述\\nSQL 注入是常见漏洞。\\n版权所有 © 2025\\n联系我们\\n## 修复方法\\n使用参数化查询。\\n'
print(TextCleaner().clean(t))
" """)
    add_para(doc, "（2）代码块与 Payload 保留验证：")
    add_code_block(doc, """python -c "
from app.rag.cleaner import TextCleaner
t = '## 代码示例\\n\\n```java\\nString sql = \\\"SELECT * FROM users WHERE id=\\\" + id;\\n```\\n\\n首页\\n## 检测载荷\\n<script>alert(1)</script>\\n'
print(TextCleaner().clean(t))
" """)
    add_h3(doc, "预期结果与验证点")
    add_table(
        doc,
        ["验证点", "预期"],
        [
            ["噪声删除", "首页 / 版权所有 / 联系我们 等行被删除，知识标题与正文保留"],
            ["代码块保留", "```java 围栏与代码行原样保留，即使相邻存在噪声行"],
            ["Payload 保留", "<script>alert(1)</script> 完整保留（默认不剥离 HTML 标签）"],
            ["CWE 编号", "CWE-89 等编号不被删除"],
            ["重复行", "连续重复的行只保留一次"],
            ["空白规范", "连续空行最多保留 1 个；行内多空格被压缩；行首缩进（列表/引用）保留"],
            ["metadata", "清洗后 Document 的 metadata 与清洗前一致"],
        ],
        col_widths=[4.0, 11.8],
    )
    add_para(doc, "", space_after=4)

    # ---- 4.3 文本分割 ----
    add_h2(doc, "4.3 文本分割测试")
    add_h3(doc, "测试目的")
    add_para(doc, "验证 Chunk 切片能生成多个语义完整的 Chunk：标题 / 描述 / 代码示例尽量不拆散，代码块不被从中切开。")
    add_h3(doc, "操作步骤")
    add_para(doc, "（1）用较小 chunk_size 观察多 Chunk 效果（便于肉眼验证语义完整）：")
    add_code_block(doc, """python -c "
from app.rag.loader import load_documents
from app.rag.splitter import SemanticRecursiveTextSplitter
d = load_documents('data/raw/owasp_cwe/sql_injection.md')[0]
s = SemanticRecursiveTextSplitter(chunk_size=300, chunk_overlap=50)
for c in s.split_document(d):
    print('---', c.metadata['chunk_index'], '| token:', c.metadata['token_count'], '| section:', c.metadata.get('section'))
    print(c.page_content[:60].replace(chr(10),' / '))
" """)
    add_para(doc, "（2）Chunk 参数实验（论文实验一预演）：依次用 chunk_size=500 / 800 / 1000 / 1500 运行管线，观察 summary.json 中的 chunk 数、token 均值等指标：")
    add_code_block(doc, """scripts\\data_pipeline\\run_pipeline.cmd --chunk-size 500
scripts\\data_pipeline\\run_pipeline.cmd --chunk-size 800
scripts\\data_pipeline\\run_pipeline.cmd --chunk-size 1000
scripts\\data_pipeline\\run_pipeline.cmd --chunk-size 1500""")
    add_h3(doc, "预期结果与验证点")
    add_table(
        doc,
        ["验证点", "预期"],
        [
            ["多 Chunk 生成", "sql_injection.md 在 chunk_size=300 时生成 2 个以上 Chunk"],
            ["语义完整", "漏洞描述 / 修复方法等段落尽量落在同一 Chunk；代码块（``` 围栏成对）不被从中切开"],
            ["token 范围", "Chunk token 数不超过 chunk_size 的约 1.5 倍（超大代码块除外，其按行二次切分）"],
            ["内容不丢失", "所有 Chunk 拼接后仍包含原文关键信息（CWE-89、PreparedStatement 等）"],
            ["metadata", "每个 Chunk 含 chunk_index / chunk_id（唯一且稳定）/ token_count / section，并保留 source / document_name / category / cwe_id 等原 metadata"],
            ["overlap", "相邻 Chunk 之间存在重叠文本（参数为 chunk_overlap）"],
        ],
        col_widths=[4.0, 11.8],
    )
    add_para(doc, "", space_after=4)

    # ---- 4.4 Metadata ----
    add_h2(doc, "4.4 Metadata 测试")
    add_h3(doc, "测试目的")
    add_para(doc, "验证知识引用提取（CWE / CVE / OWASP）与 ChromaDB 兼容的规范化。")
    add_h3(doc, "操作步骤")
    add_code_block(doc, """python -c "
from app.rag.metadata import extract_knowledge_metadata, normalize_metadata
print(extract_knowledge_metadata('CWE-89 与 OWASP A03:2021 相关，参考 CVE-2024-12345'))
print(normalize_metadata({'a': None, 'b': ' ok ', 'tags': ['x','y'], 'n': 3}))
" """)
    add_h3(doc, "预期结果与验证点")
    add_table(
        doc,
        ["验证点", "预期"],
        [
            ["CWE 提取", "识别 CWE-89（含 cwe-89 小写形式），输出列表 [CWE-89]"],
            ["CVE 提取", "识别 CVE-2024-12345（含下划线分隔 CVE-2021_44228），输出带 CVE- 前缀"],
            ["OWASP 提取", "识别 A03:2021 等编号"],
            ["规范化", "None 与空字符串剔除；字符串去除首尾空白；列表保留；字典等复杂对象降级为字符串"],
            ["ChromaDB 兼容", "规范化后所有值类型均为 str / int / float / bool / 字符串列表"],
        ],
        col_widths=[4.0, 11.8],
    )
    add_para(doc, "", space_after=4)

    # ---- 4.5 完整管线 ----
    add_h2(doc, "4.5 完整管线测试（raw → processed）")
    add_h3(doc, "测试目的")
    add_para(doc, "验证从原始知识数据到 Chunk 产物的完整链路，产物可供后续 Embedding / ChromaDB 阶段直接使用。")
    add_h3(doc, "操作步骤")
    add_para(doc, "（1）一键运行管线：")
    add_code_block(doc, """scripts\\data_pipeline\\run_pipeline.cmd
# 或 python scripts/data_pipeline/build_knowledge_base.py""")
    add_para(doc, "（2）检查处理产物：")
    add_code_block(doc, """type backend\\data\\processed\\summary.json          # 统计信息（Windows）
python -c "import json; [print(json.loads(l)['metadata']['chunk_id'], json.loads(l)['metadata']['token_count']) for l in open('backend/data/processed/chunks.jsonl',encoding='utf-8')]"  # 抽查 Chunk""")
    add_h3(doc, "预期结果与验证点")
    add_table(
        doc,
        ["验证点", "预期"],
        [
            ["管线执行", "日志依次显示 1/4 加载 → 2/4 清洗 → 3/4 Metadata 增强 → 4/4 分割，退出码 0"],
            ["产物文件", "生成 backend/data/processed/chunks.jsonl（JSON Lines）与 summary.json"],
            ["统计合理", "当前 16 个文档约生成 21 个 Chunk；token 范围约 55~795，均值约 490"],
            ["Chunk 结构", "每行含 page_content 与 metadata；metadata 含 source / document_name / file_type / category / chunk_id / chunk_index / token_count，知识文档还含 cwe_id / owasp_id / section"],
            ["可复现", "重复运行结果稳定（chunk_id 基于来源路径哈希生成）"],
            ["无副作用", "原始数据 backend/data/raw/ 不被修改"],
        ],
        col_widths=[4.0, 11.8],
    )
    add_para(doc, "", space_after=4)

    # ================= 5. 测试验证清单 =================
    add_h1(doc, "5. 测试验证清单")
    add_para(doc, "开发 / 验收时逐项勾选：")
    add_table(
        doc,
        ["#", "检查项", "通过标准", "结果"],
        [
            ["1", "自动化测试", "pytest 全部通过（50 passed）", "□"],
            ["2", "文档加载", "5 种格式均可加载，metadata 完整（source/document_name/file_type/category）", "□"],
            ["3", "知识数据可读", "data/raw 下 16 个文件全部加载成功，失败数 0", "□"],
            ["4", "文本清洗", "噪声被删除，代码示例 / Payload / CWE 编号完整保留", "□"],
            ["5", "文本分割", "生成多个 Chunk，代码块不被切开，token 在合理范围", "□"],
            ["6", "Metadata", "CWE/CVE/OWASP 可提取，规范化值类型与 ChromaDB 兼容", "□"],
            ["7", "完整管线", "chunks.jsonl + summary.json 生成，统计合理，可复现", "□"],
            ["8", "来源溯源", "每个 Chunk 的 metadata.source 可追溯到原始文件", "□"],
            ["9", "范围检查", "未开发 Embedding / ChromaDB / Retriever / Agent 等阶段外功能", "□"],
        ],
        col_widths=[1.0, 3.4, 9.2, 2.2],
    )
    add_para(doc, "", space_after=4)

    # ================= 6. 常见问题与排查 =================
    add_h1(doc, "6. 常见问题与排查")
    add_table(
        doc,
        ["现象", "原因", "处理方法"],
        [
            ["ModuleNotFoundError: No module named 'app'", "PYTHONPATH 未包含 backend 目录", "在 backend 目录下运行，或设置 PYTHONPATH=backend；使用 scripts\\launch\\run_backend.cmd 等一键脚本"],
            ["ModuleNotFoundError: pypdf / bs4 等", "依赖未安装", "pip install -r requirements.txt；本环境使用 .deps 时确认 PYTHONPATH 含 .deps"],
            ["控制台中文乱码", "Windows 控制台默认 GBK，而日志/输出为 UTF-8", "改用脚本文件重定向输出后查看，或 chcp 65001 切换代码页"],
            ["pip 安装报 WinError 5", "安装到用户目录被权限拦截（沙箱环境）", "pip install --target <项目>\\.deps，通过 PYTHONPATH 使用"],
            ["Chunk 偏大/偏小", "chunk_size 参数不适合当前数据", "调整 --chunk-size / --chunk-overlap；论文实验建议比较 500/800/1000/1500"],
            ["PDF 无法加载", "PDF 加密或文本层缺失", "确认文件未加密；扫描版 PDF 需 OCR 预处理（后续阶段）"],
            ["清洗删除了代码", "噪声模式配置过宽", "检查 config.py 中 noise_line_patterns，代码块保护机制不应对围栏内内容生效；可单独验证 TextCleaner().clean() 结果"],
        ],
        col_widths=[5.2, 4.6, 6.0],
    )
    add_para(doc, "", space_after=6)

    add_para(
        doc,
        "附：本指南对应的测试数据与产物均可复现；如模块行为调整，请同步更新自动化测试与本文档。",
        size=9.5,
        color=(0x88, 0x88, 0x88),
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(f"文档已生成: {out} ({out.stat().st_size} bytes)")
    sys.exit(0)
